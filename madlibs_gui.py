# MADLIBS GUI
import PySimpleGUI as sg
import os
from docx import Document
from docx.shared import Pt , RGBColor


#To see all possible themes of the PySimpleGUI, do : sg.preview_all_look_and_feel_themes()
sg.theme("LightBlue2")

# First input of the margins() function is for breadth and second is for height.
# layout[] handles what is displayed to the user. That is, the GUI components
# read() returns any events triggered in the Window() function
enter = sg.Button("Enter", size=(5, 1), button_color="Black", font='Any 15')
layout = [  [sg.Text("Hello, cohorts! Welcome to our MADLIBS Game!", size=(45, 3), font="arial 20 bold", text_color="Black")],
            [sg.Text("How do u feel ? (emotion) : "), sg.InputText()],
            [sg.Text("What is your favourite color ? "), sg.InputText()],
            [sg.Text("What is your favourite weather ? "), sg.InputText()],
            [sg.Text("Enter clothing brand name : "), sg.InputText()],
            [sg.Text("Describe something with 4 words separated by ',' (rigid , active , bright ...) : "), sg.InputText()],
            [sg.Text("Enter 2 Nouns seperated by ',' : "), sg.InputText()],
            [sg.Text("Enter a verb/action : "), sg.InputText()],
            [sg.Text("Enter footwear name : "), sg.InputText()],
            [sg.Text("Enter vehicle name : "), sg.InputText()],
            [sg.Text("Enter your town name : "), sg.InputText()],
            [sg.Text("Enter a direction (north , west ... ) : "), sg.InputText()],
            [sg.Text("What is your favourite food ? "), sg.InputText()],
            [sg.Text("What is the animal you'll like to take care of ? "), sg.InputText()],
            [sg.Text("What is your vibe for this story? ", text_color="Black"), sg.InputText()],
            [sg.Text("Give me a name for a document? ", text_color="Black"), sg.InputText()],
            [enter]]

# create the Window. Window() function loads a GUI window.
window = sg.Window("POLLY MADLIBS", layout)

# checking if user put enough inputs
def trustless_input(LIST, number):
    for i in range(number - len(LIST)):
        LIST.append(LIST[i])  # making up for the inputs!

# A graphical user interface needs to run inside a loop and wait for the user to do something.
# We create an event loop, which can only end when we satisfy the set condition.
while True:
    event, values = window.read()
    # End program if the user closes the window
    if event == sg.WINDOW_CLOSED:
        break
    elif event == "Enter":
        # get and assign the inputs to variables
        emotion = values[0]
        color = values[1]
        weather = values[2]
        cloth = values[3]
        adj = values[4]
        adj = adj.split(',')
        noun = values[5]
        noun = noun.split(',')
        verb = values[6]
        footwear = values[7]
        vehicle = values[8]
        town = values[9]
        direction = values[10]
        food = values[11]
        animal = values[12]
        title=values[13]
        name=values[14]

        trustless_input(adj, 4)
        trustless_input(noun, 2)

        # initiate Document
        document = Document()
        # styling
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(16)  #

        # general Heading
        document.add_heading('Mad Libs Project Team POLLY', 0)

        # text heading

        title = 'The {} Day Ever'.format(title)
        document.add_heading(title, level=1)
        # paragraph 1
        p1 = document.add_paragraph('	One day I woke up feeling  ')
        emotion = p1.add_run(emotion)
        emotion.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(' , and I knew it was going to be a special day. The sky was ')
        color = p1.add_run(color)
        color.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(' and the weather was ')
        weather = p1.add_run(weather)
        weather.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(', so I hopped out of bed, put on my ')
        cloth = p1.add_run(cloth)
        cloth.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(' and my ')
        adj0 = p1.add_run(adj[0])
        adj0.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(' ')
        footwear = p1.add_run(footwear)
        footwear.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p1.add_run(', and I was ready to ')
        verb = p1.add_run(verb+".")
        verb.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)

        # paragraph 2
        p2 = document.add_paragraph('	Outside, I caught the first ')
        vehicle = p2.add_run(vehicle)
        vehicle.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run(', which took me straight into ')
        town = p2.add_run(town)
        town.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run('. I went ')
        direction = p2.add_run(direction)
        direction.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run(' until I came to a store selling ')
        adj1 = p2.add_run(adj[1])
        adj1.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        noun0 = p2.add_run(noun[0])
        noun0.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run(', where I bought the perfect ')
        adj2 = p2.add_run(adj[2])
        adj2.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run(' ')
        noun1 = p2.add_run(noun[1])
        noun1.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p2.add_run('!')

        # paragraph 3
        p3 = document.add_paragraph('	Next, I treated myself to a ')
        food = p3.add_run(food)
        food.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p3.add_run(' snack at a restaurant. ')

        # paragraph 4
        p4 = document.add_paragraph('	Finally, I went back home. I fed the ')
        animals = p4.add_run(animal)
        animals.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p4.add_run(', then sat down on the furniture , and thought, "What a ')
        adj3 = p4.add_run(adj[3])
        adj3.font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
        p4.add_run(' day!" ')

        # save file and run it
        name = name + ".docx"
        document.save(name)
        start = "start {}".format(name)
        os.system(start)



